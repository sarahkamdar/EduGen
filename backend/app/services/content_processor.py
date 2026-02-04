import os
import re
from pathlib import Path
from typing import Optional, Callable, Awaitable
from fastapi import UploadFile, HTTPException
from app.services.audio_extractor import extract_audio
from app.services.transcription import transcribe_audio

async def normalize_video(file: UploadFile, progress_callback: Optional[Callable[[str, str, int], Awaitable[None]]] = None) -> str:
    temp_dir = Path("temp")
    temp_dir.mkdir(exist_ok=True)
    
    if progress_callback:
        await progress_callback("upload", "Saving video file...", 10)
    
    file_path = temp_dir / file.filename
    with open(file_path, "wb") as f:
        f.write(await file.read())
    
    if progress_callback:
        await progress_callback("extract", "Extracting audio from video...", 30)
    
    audio_path = extract_audio(str(file_path))
    os.remove(file_path)
    
    if progress_callback:
        await progress_callback("transcribe", "Transcribing audio to text...", 60)
    
    transcript = transcribe_audio(audio_path)
    os.remove(audio_path)
    
    if progress_callback:
        await progress_callback("finalize", "Finalizing content...", 90)
    
    return transcript

async def normalize_youtube(youtube_url: str, progress_callback: Optional[Callable[[str, str, int], Awaitable[None]]] = None) -> str:
    try:
        import yt_dlp
        
        temp_dir = Path("temp")
        temp_dir.mkdir(exist_ok=True)
        
        if progress_callback:
            await progress_callback("upload", "Connecting to YouTube...", 10)
        
        audio_file = temp_dir / "youtube_audio.mp3"
        
        ydl_opts = {
            'format': 'bestaudio/best',
            'outtmpl': str(audio_file.with_suffix('')),
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
            }],
            'quiet': False,
            'no_warnings': False,
            'extractor_args': {
                'youtube': {
                    'player_client': ['android', 'web'],
                    'skip': ['dash', 'hls']
                }
            },
            'http_headers': {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
            }
        }
        
        if progress_callback:
            await progress_callback("extract", "Downloading audio from YouTube...", 30)
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([youtube_url])
        
        if progress_callback:
            await progress_callback("transcribe", "Transcribing audio to text...", 60)
        
        transcript = transcribe_audio(str(audio_file))
        os.remove(audio_file)
        
        if progress_callback:
            await progress_callback("finalize", "Finalizing content...", 90)
        
        return transcript
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"YouTube download failed: {str(e)}")

async def normalize_pdf(file: UploadFile, progress_callback: Optional[Callable[[str, str, int], Awaitable[None]]] = None) -> str:
    try:
        from pypdf import PdfReader
        
        temp_dir = Path("temp")
        temp_dir.mkdir(exist_ok=True)
        
        if progress_callback:
            await progress_callback("upload", "Saving PDF file...", 10)
        
        file_path = temp_dir / file.filename
        with open(file_path, "wb") as f:
            f.write(await file.read())
        
        if progress_callback:
            await progress_callback("extract", "Extracting text from PDF...", 40)
        
        text = ""
        try:
            pdf_reader = PdfReader(file_path)
            total_pages = len(pdf_reader.pages)
            
            for i, page in enumerate(pdf_reader.pages):
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
                
                # Update progress per page
                if progress_callback and total_pages > 0:
                    page_progress = 40 + int((i + 1) / total_pages * 40)
                    await progress_callback("extract", f"Processing page {i+1} of {total_pages}...", page_progress)
        finally:
            if os.path.exists(file_path):
                os.remove(file_path)
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No readable text found in PDF. It may be scanned or image-based.")
        
        if progress_callback:
            await progress_callback("finalize", "Finalizing content...", 90)
        
        return clean_text(text)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"PDF processing failed: {str(e)}")

async def normalize_word(file: UploadFile, progress_callback: Optional[Callable[[str, str, int], Awaitable[None]]] = None) -> str:
    try:
        from docx import Document
        
        temp_dir = Path("temp")
        temp_dir.mkdir(exist_ok=True)
        
        if progress_callback:
            await progress_callback("upload", "Saving Word document...", 10)
        
        file_path = temp_dir / file.filename
        with open(file_path, "wb") as f:
            f.write(await file.read())
        
        if progress_callback:
            await progress_callback("extract", "Extracting text from document...", 40)
        
        text = ""
        try:
            doc = Document(file_path)
            total_paragraphs = len(doc.paragraphs)
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                
                # Update progress
                if progress_callback and total_paragraphs > 0:
                    para_progress = 40 + int((i + 1) / total_paragraphs * 30)
                    await progress_callback("extract", f"Processing paragraph {i+1} of {total_paragraphs}...", para_progress)
            
            if progress_callback:
                await progress_callback("extract", "Extracting tables...", 75)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
        finally:
            if os.path.exists(file_path):
                os.remove(file_path)
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No readable text found in Word document.")
        
        if progress_callback:
            await progress_callback("finalize", "Finalizing content...", 90)
        
        return clean_text(text)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Word document processing failed: {str(e)}")

def normalize_text(text: str) -> str:
    return clean_text(text)

def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    return text

async def process_content(
    file: Optional[UploadFile] = None,
    youtube_url: Optional[str] = None,
    text: Optional[str] = None
) -> tuple[str, str]:
    input_count = sum([file is not None, youtube_url is not None, text is not None])
    
    if input_count == 0:
        raise HTTPException(status_code=400, detail="No input provided")
    
    if input_count > 1:
        raise HTTPException(status_code=400, detail="Only one input type allowed")
    
    if file:
        file_ext = Path(file.filename).suffix.lower()
        
        video_extensions = ['.mp4', '.avi', '.mov', '.mkv', '.flv', '.wmv']
        pdf_extensions = ['.pdf']
        word_extensions = ['.docx', '.doc']
        
        if file_ext in video_extensions:
            normalized_text = await normalize_video(file)
            input_type = "video"
        elif file_ext in pdf_extensions:
            normalized_text = await normalize_pdf(file)
            input_type = "pdf"
        elif file_ext in word_extensions:
            normalized_text = await normalize_word(file)
            input_type = "word"
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format")
    
    elif youtube_url:
        normalized_text = await normalize_youtube(youtube_url)
        input_type = "youtube"
    
    elif text:
        normalized_text = normalize_text(text)
        input_type = "text"
    
    return input_type, normalized_text

async def process_content_with_progress(
    file: Optional[UploadFile] = None,
    youtube_url: Optional[str] = None,
    text: Optional[str] = None,
    progress_callback: Optional[Callable[[str, str, int], Awaitable[None]]] = None
) -> tuple[str, str]:
    """Process content with real-time progress updates"""
    
    input_count = sum([file is not None, youtube_url is not None, text is not None])
    
    if input_count == 0:
        raise HTTPException(status_code=400, detail="No input provided")
    
    if input_count > 1:
        raise HTTPException(status_code=400, detail="Only one input type allowed")
    
    if progress_callback:
        await progress_callback("start", "Starting content processing...", 5)
    
    if file:
        file_ext = Path(file.filename).suffix.lower()
        
        video_extensions = ['.mp4', '.avi', '.mov', '.mkv', '.flv', '.wmv']
        pdf_extensions = ['.pdf']
        word_extensions = ['.docx', '.doc']
        
        if file_ext in video_extensions:
            normalized_text = await normalize_video(file, progress_callback)
            input_type = "video"
        elif file_ext in pdf_extensions:
            normalized_text = await normalize_pdf(file, progress_callback)
            input_type = "pdf"
        elif file_ext in word_extensions:
            normalized_text = await normalize_word(file, progress_callback)
            input_type = "word"
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format")
    
    elif youtube_url:
        normalized_text = await normalize_youtube(youtube_url, progress_callback)
        input_type = "youtube"
    
    elif text:
        if progress_callback:
            await progress_callback("upload", "Processing text input...", 30)
            await progress_callback("extract", "Normalizing text...", 60)
            await progress_callback("finalize", "Finalizing content...", 90)
        normalized_text = normalize_text(text)
        input_type = "text"
    
    return input_type, normalized_text
